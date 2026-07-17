"""
Generate a conclusion document (Заключение об открытом публиковании)
by filling placeholders in the DOCX template.

Placeholders in template:
  [название_статьи]   → article title
  [ФИО_авторов]       → abbreviated author names (Иванов И.И., Петров П.П.)
  [окончание_автор]   → suffix: "а" for 1 author, "ов" for 2+
  [месяц_загрузки]    → month name in genitive case (ru)
  [год_загрузки]      → 4-digit year
  [главный_автор]     → lead author abbreviated name (or first author)
"""

import io
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document

BUILTIN_TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "conclusion_template.docx"

MONTHS_GENITIVE = {
    1:  "января",
    2:  "февраля",
    3:  "марта",
    4:  "апреля",
    5:  "мая",
    6:  "июня",
    7:  "июля",
    8:  "августа",
    9:  "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}


def abbreviate_name(full_name: str) -> str:
    """
    Convert full name to abbreviated form.
    'Лаукарт Михаил Сергеевич' -> 'Лаукарт М.С.'
    'Иванов Иван' -> 'Иванов И.'
    """
    parts = full_name.strip().split()
    if not parts:
        return full_name
    surname = parts[0]
    initials = "".join(p[0].upper() + "." for p in parts[1:] if p)
    return f"{surname} {initials}".strip()


def _replace_in_paragraph(para, replacements: dict) -> None:
    """
    Merge all runs of a paragraph into a single text string,
    apply replacements, then redistribute back into runs
    preserving the formatting of the first run.
    """
    full_text = "".join(r.text for r in para.runs)
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, value)

    if new_text == full_text:
        return

    for run in para.runs:
        run.text = ""

    if para.runs:
        para.runs[0].text = new_text


def generate_conclusion(
    article_title: str,
    author_names: List[str],
    upload_date: datetime,
    lead_author_name: Optional[str] = None,
    template_path=None,
) -> bytes:
    """
    Fill the conclusion template and return the resulting DOCX as bytes.
    Uses template_path if provided, otherwise falls back to built-in template.
    """
    doc = Document(str(template_path or BUILTIN_TEMPLATE_PATH))

    month_name = MONTHS_GENITIVE[upload_date.month]
    year_str = str(upload_date.year)

    # Abbreviated author names: Лаукарт М.С., Яссер М.В.
    abbreviated = [abbreviate_name(n) for n in author_names] if author_names else []
    authors_str = ", ".join(abbreviated) if abbreviated else "—"

    # Suffix: "а" for 1 author, "ов" for 2+
    author_suffix = "а" if len(author_names) == 1 else "ов"

    # Lead author
    if lead_author_name:
        lead_str = abbreviate_name(lead_author_name)
    elif author_names:
        lead_str = abbreviate_name(author_names[0])
    else:
        lead_str = "—"

    replacements = {
        # New placeholders
        "[месяц_загрузки]": month_name,
        "[год_загрузки]": year_str,
        "[название_статьи]": article_title,
        "[ФИО_авторов]": authors_str,
        "[окончание_автор]": author_suffix,
        "[главный_автор]": lead_str,
        # Legacy placeholders (backward compatibility)
        "[Месяц_загрузки]": month_name,
        "[Год_загрузки]": year_str,
        "[имя статьи]": article_title,
        "[а|ов]": author_suffix,
    }

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
